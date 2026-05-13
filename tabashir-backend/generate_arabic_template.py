from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

def make_rtl(paragraph):
    # Ensure paragraph is aligned to the right
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Inject w:bidi tag correctly so Word respects it
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.insert_element_before(bidi,
        'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )

def make_rtl_run(run):
    # Inject w:rtl tag into the run properties so the text flows right-to-left
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.insert_element_before(rtl,
        'w:cs', 'w:emb', 'w:sz', 'w:szCs', 'w:effect', 'w:bdr', 'w:shd',
        'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang',
        'w:eastAsianLayout', 'w:specVanish', 'w:oMath'
    )

def add_rtl_run(paragraph, text):
    run = paragraph.add_run(text)
    make_rtl_run(run)
    return run

doc = Document()

# Set Margins (0.75 inch ATS friendly)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Set Default Font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# --- HEADER ---
name_p = doc.add_paragraph()
make_rtl(name_p)
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = add_rtl_run(name_p, "{{ header.name }}")
name_run.bold = True
name_run.font.size = Pt(16)

contact_p = doc.add_paragraph()
make_rtl(contact_p)
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = add_rtl_run(contact_p, "{{ contact_info }}")
contact_run.font.size = Pt(10)

# --- SUMMARY ---
doc.add_paragraph("{% if objective %}")
p = doc.add_paragraph()
make_rtl(p)
run = add_rtl_run(p, "الملخص المهني")
run.bold = True
run.font.size = Pt(12)
add_horizontal_line(p)

obj_p = doc.add_paragraph()
make_rtl(obj_p)
add_rtl_run(obj_p, "{{ objective }}")
doc.add_paragraph("{% endif %}")

# --- WORK EXPERIENCE ---
doc.add_paragraph("{% if work_list %}")
p = doc.add_paragraph()
make_rtl(p)
run = add_rtl_run(p, "الخبرة المهنية")
run.bold = True
run.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for work in work_list %}")
work_p = doc.add_paragraph()
make_rtl(work_p)
add_rtl_run(work_p, "{{ work.position }}").bold = True
add_rtl_run(work_p, "{% if work.company %} | {{ work.company }}{% endif %}").italic = True
add_rtl_run(work_p, "{% if work.location %} | {{ work.location }}{% endif %}{% if work.date %} | {{ work.date }}{% endif %}")

doc.add_paragraph("{% for detail in work.details %}")
detail_p = doc.add_paragraph()
make_rtl(detail_p)
add_rtl_run(detail_p, "•  {{ detail }}")
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- LEADERSHIP EXPERIENCE ---
doc.add_paragraph("{% if lship_list %}")
p = doc.add_paragraph()
make_rtl(p)
run = add_rtl_run(p, "الخبرة القيادية")
run.bold = True
run.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for lship in lship_list %}")
lship_p = doc.add_paragraph()
make_rtl(lship_p)
add_rtl_run(lship_p, "{{ lship.position }}").bold = True
add_rtl_run(lship_p, "{% if lship.company %} | {{ lship.company }}{% endif %}").italic = True
add_rtl_run(lship_p, "{% if lship.location %} | {{ lship.location }}{% endif %}{% if lship.date %} | {{ lship.date }}{% endif %}")

doc.add_paragraph("{% for detail in lship.details %}")
detail_p = doc.add_paragraph()
make_rtl(detail_p)
add_rtl_run(detail_p, "•  {{ detail }}")
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- EDUCATION ---
doc.add_paragraph("{% if education_list %}")
p = doc.add_paragraph()
make_rtl(p)
run = add_rtl_run(p, "التعليم")
run.bold = True
run.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for edu in education_list %}")
edu_p = doc.add_paragraph()
make_rtl(edu_p)
add_rtl_run(edu_p, "{{ edu.degree }}{% if edu.major %} في {{ edu.major }}{% endif %}").bold = True
add_rtl_run(edu_p, "{% if edu.university %} | {{ edu.university }}{% endif %}").italic = True
add_rtl_run(edu_p, "{% if edu.location %} | {{ edu.location }}{% endif %}{% if edu.date %} | {{ edu.date }}{% endif %}")

doc.add_paragraph("{% if edu.gpa and edu.gpa != 'None' %}")
gpa_p = doc.add_paragraph()
make_rtl(gpa_p)
add_rtl_run(gpa_p, "المعدل: {{ edu.gpa }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% for detail in edu.details %}")
detail_p = doc.add_paragraph()
make_rtl(detail_p)
add_rtl_run(detail_p, "•  {{ detail }}")
doc.add_paragraph("{% endfor %}")

doc.add_paragraph("{% for course in edu.coursework %}")
course_p = doc.add_paragraph()
make_rtl(course_p)
add_rtl_run(course_p, "•  الدورات: {{ course }}")
doc.add_paragraph("{% endfor %}")

doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- PROJECTS ---
doc.add_paragraph("{% if project_list %}")
p = doc.add_paragraph()
make_rtl(p)
run = add_rtl_run(p, "المشاريع")
run.bold = True
run.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for project in project_list %}")
proj_p = doc.add_paragraph()
make_rtl(proj_p)
add_rtl_run(proj_p, "{{ project.title }}").bold = True
add_rtl_run(proj_p, "{% if project.position %} | {{ project.position }}{% endif %}").italic = True
add_rtl_run(proj_p, "{% if project.location %} | {{ project.location }}{% endif %}{% if project.date %} | {{ project.date }}{% endif %}")

doc.add_paragraph("{% for detail in project.details %}")
detail_p = doc.add_paragraph()
make_rtl(detail_p)
add_rtl_run(detail_p, "•  {{ detail }}")
doc.add_paragraph("{% endfor %}")

doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- SKILLS & LANGUAGES ---
doc.add_paragraph("{% if skills.skillset or skills.softskills or languages %}")
p = doc.add_paragraph()
make_rtl(p)
run = add_rtl_run(p, "المهارات والكفاءات")
run.bold = True
run.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% if skills.skillset %}")
sk_p = doc.add_paragraph()
make_rtl(sk_p)
add_rtl_run(sk_p, "المهارات التقنية: ").bold = True
add_rtl_run(sk_p, "{{ skills.skillset|join('، ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% if skills.softskills %}")
ss_p = doc.add_paragraph()
make_rtl(ss_p)
add_rtl_run(ss_p, "الكفاءات الأساسية: ").bold = True
add_rtl_run(ss_p, "{{ skills.softskills|join('، ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% if languages %}")
lang_p = doc.add_paragraph()
make_rtl(lang_p)
add_rtl_run(lang_p, "اللغات: ").bold = True
add_rtl_run(lang_p, "{{ languages|join('، ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% if skills.training %}")
tr_p = doc.add_paragraph()
make_rtl(tr_p)
add_rtl_run(tr_p, "التدريب والشهادات: ").bold = True
add_rtl_run(tr_p, "{{ skills.training|join('، ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% endif %}")

doc.save('tabashir-backend/templates/Arabic Docxtpl Compatible CV Template.docx')
print("Arabic template generated successfully.")