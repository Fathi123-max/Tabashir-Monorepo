from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

doc = Document()

# Set Margins (0.75 inch ATS friendly)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Set Default Font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- HEADER ---
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name_p.add_run("{{ header.name }}")
name_run.bold = True
name_run.font.size = Pt(16)

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact_p.add_run("{{ contact_info }}")
contact_run.font.size = Pt(10)

# --- SUMMARY ---
doc.add_paragraph("{% if objective %}")
p = doc.add_paragraph("PROFESSIONAL SUMMARY")
p.style.font.bold = True
p.style.font.size = Pt(12)
add_horizontal_line(p)
doc.add_paragraph("{{ objective }}")
doc.add_paragraph("{% endif %}")

# --- WORK EXPERIENCE ---
doc.add_paragraph("{% if work_list %}")
p = doc.add_paragraph("WORK EXPERIENCE")
p.style.font.bold = True
p.style.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for work in work_list %}")
work_p = doc.add_paragraph()
work_p.add_run("{{ work.position }}").bold = True
work_p.add_run("{% if work.company %} | {{ work.company }}{% endif %}").italic = True
work_p.add_run("{% if work.location %} | {{ work.location }}{% endif %}{% if work.date %} | {{ work.date }}{% endif %}")

doc.add_paragraph("{% for detail in work.details %}")
# Normal bullet points without extra characters
detail_p = doc.add_paragraph("{{ detail }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- LEADERSHIP EXPERIENCE ---
doc.add_paragraph("{% if lship_list %}")
p = doc.add_paragraph("LEADERSHIP EXPERIENCE")
p.style.font.bold = True
p.style.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for lship in lship_list %}")
lship_p = doc.add_paragraph()
lship_p.add_run("{{ lship.position }}").bold = True
lship_p.add_run("{% if lship.company %} | {{ lship.company }}{% endif %}").italic = True
lship_p.add_run("{% if lship.location %} | {{ lship.location }}{% endif %}{% if lship.date %} | {{ lship.date }}{% endif %}")

doc.add_paragraph("{% for detail in lship.details %}")
doc.add_paragraph("{{ detail }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")


# --- EDUCATION ---
doc.add_paragraph("{% if education_list %}")
p = doc.add_paragraph("EDUCATION")
p.style.font.bold = True
p.style.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for edu in education_list %}")
edu_p = doc.add_paragraph()
edu_p.add_run("{{ edu.degree }}{% if edu.major %} in {{ edu.major }}{% endif %}").bold = True
edu_p.add_run("{% if edu.university %} | {{ edu.university }}{% endif %}").italic = True
edu_p.add_run("{% if edu.location %} | {{ edu.location }}{% endif %}{% if edu.date %} | {{ edu.date }}{% endif %}")

doc.add_paragraph("{% if edu.gpa and edu.gpa != 'None' %}")
doc.add_paragraph("GPA: {{ edu.gpa }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% for detail in edu.details %}")
doc.add_paragraph("{{ detail }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")

doc.add_paragraph("{% for course in edu.coursework %}")
doc.add_paragraph("Coursework: {{ course }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")

doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- PROJECTS ---
doc.add_paragraph("{% if project_list %}")
p = doc.add_paragraph("PROJECTS")
p.style.font.bold = True
p.style.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% for project in project_list %}")
proj_p = doc.add_paragraph()
proj_p.add_run("{{ project.title }}").bold = True
proj_p.add_run("{% if project.position %} | {{ project.position }}{% endif %}").italic = True
proj_p.add_run("{% if project.location %} | {{ project.location }}{% endif %}{% if project.date %} | {{ project.date }}{% endif %}")

doc.add_paragraph("{% for detail in project.details %}")
doc.add_paragraph("{{ detail }}", style='List Bullet')
doc.add_paragraph("{% endfor %}")

doc.add_paragraph("{% endfor %}")
doc.add_paragraph("{% endif %}")

# --- SKILLS & LANGUAGES ---
doc.add_paragraph("{% if skills.skillset or skills.softskills or languages %}")
p = doc.add_paragraph("SKILLS & COMPETENCIES")
p.style.font.bold = True
p.style.font.size = Pt(12)
add_horizontal_line(p)

doc.add_paragraph("{% if skills.skillset %}")
sk_p = doc.add_paragraph()
sk_p.add_run("Technical Skills: ").bold = True
sk_p.add_run("{{ skills.skillset|join(', ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% if skills.softskills %}")
ss_p = doc.add_paragraph()
ss_p.add_run("Core Competencies: ").bold = True
ss_p.add_run("{{ skills.softskills|join(', ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% if languages %}")
lang_p = doc.add_paragraph()
lang_p.add_run("Languages: ").bold = True
lang_p.add_run("{{ languages|join(', ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% if skills.training %}")
tr_p = doc.add_paragraph()
tr_p.add_run("Training & Certifications: ").bold = True
tr_p.add_run("{{ skills.training|join(', ') }}")
doc.add_paragraph("{% endif %}")

doc.add_paragraph("{% endif %}")

doc.save('tabashir-backend/templates/Docxtpl Compatible CV Template.docx')
print("Templates generated successfully.")
