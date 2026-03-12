import openai
from docx import Document
from docx.shared import Pt, RGBColor

from app import Config
from app.services.cv_processor import get_openai_client


def extract_text_with_format(file_path):
    """
    Extracts text and formatting from a DOCX file.
    """
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        formatted_text = []
        for run in paragraph.runs:
            formatted_text.append({
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "font_color": run.font.color.rgb if run.font.color else None
            })
        content.append(formatted_text)
    return content


def translate_text_to_arabic(text):
    """
    Translates text to Arabic using OpenAI's GPT API.
    """
    client, model = get_openai_client()
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "Translate this to Arabic."},
            {"role": "user", "content": text}
        ]
    )
    return response.choices[0].message.content


def save_translated_docx(translated_text, format_template, output_path):
    """
    Saves translated a text into a DOCX file while preserving the formatting from the original.
    """
    doc = Document()
    translated_lines = translated_text.split("\n")

    for i, original_paragraph in enumerate(format_template):
        if i < len(translated_lines):
            paragraph = doc.add_paragraph()
            for run_format in original_paragraph:
                run = paragraph.add_run(translated_lines[i])
                run.bold = run_format["bold"]
                run.italic = run_format["italic"]
                run.underline = run_format["underline"]
                if run_format["font_name"]:
                    run.font.name = run_format["font_name"]
                if run_format["font_size"]:
                    run.font.size = Pt(run_format["font_size"])
                if run_format["font_color"]:
                    run.font.color.rgb = RGBColor(*run_format["font_color"])
            # One line per paragraph
    doc.save(output_path)


def translate_docx_to_arabic(input_path, output_path):
    """
    calls related methods to translate a DOCX file to Arabic.
    """
    original_format = extract_text_with_format(input_path)
    content = "\n".join([run["text"] for para in original_format for run in para if run["text"]])
    translated = translate_text_to_arabic(content)
    save_translated_docx(translated, original_format, output_path)
