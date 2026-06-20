from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def insert_bidi(pPr):
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.insert_element_before(bidi,
        'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )

def make_rtl_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    insert_bidi(pPr)

def make_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.insert_element_before(rtl,
        'w:cs', 'w:emb', 'w:sz', 'w:szCs', 'w:effect', 'w:bdr', 'w:shd',
        'w:fitText', 'w:vertAlign', 'w:rtl', 'w:cs', 'w:em', 'w:lang',
        'w:eastAsianLayout', 'w:specVanish', 'w:oMath'
    )

doc = Document()
p = doc.add_paragraph()
make_rtl_paragraph(p)
run = p.add_run("هذا نص تجريبي باللغة العربية")
make_rtl_run(run)

import os
dir_path = os.path.dirname(os.path.realpath(__file__))
doc.save(os.path.join(dir_path, "test_rtl_output.docx"))
print("RTL test document generated.")
